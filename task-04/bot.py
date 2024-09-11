from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, CallbackQueryHandler, filters
import requests
import csv
import os
from docx import Document

GOOGLE_BOOKS_API_KEY = 'AIzaSyBvtSgcggs81eybB8i6IOUHMdVkrZv6Xtw'

# Store reading list in-memory for simplicity
reading_list = []
user_state = {}  # To track the user's current action (book search, preview, add, delete)

async def start(update: Update, context):
    await update.message.reply_text("Welcome to PagePal! Use /help to see available commands.")

async def help_command(update: Update, context):
    commands = "/start - Welcome message\n" \
               "/book - Enter a genre to get book recommendations\n" \
               "/preview - Enter a book name to get its preview link\n" \
               "/list - Manage your reading list\n" \
               "/reading_list - View, Add, or Delete books from your list"
    await update.message.reply_text(commands)

async def book_command(update: Update, context):
    user_state[update.effective_user.id] = 'book'
    await update.message.reply_text("Please enter a genre to search for books.")

async def preview_command(update: Update, context):
    user_state[update.effective_user.id] = 'preview'
    await update.message.reply_text("Please enter the book name for which you need a preview link.")

async def handle_genre_search(update: Update, context):
    genre = update.message.text
    books = search_books_by_genre(genre)
    if books:
        csv_filename = f"{genre}_books.csv"
        with open(csv_filename, mode='w', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            writer.writerow(["Title", "Author", "Description", "Published Year", "Language", "Preview Link"])
            for book in books:
                writer.writerow([book['title'], book['author'], book['description'], book['published_year'],
                                 book['language'], book['preview_link']])

        with open(csv_filename, 'rb') as file:
            await update.message.reply_document(document=file, filename=csv_filename)
        os.remove(csv_filename)  # Clean up
    else:
        await update.message.reply_text("No books found for this genre.")

async def handle_preview(update: Update, context):
    book_name = update.message.text
    preview_link = get_book_preview(book_name)
    if preview_link:
        await update.message.reply_text(f"Preview link for '{book_name}': {preview_link}")
    else:
        await update.message.reply_text(f"No preview available for '{book_name}'.")

async def reading_list_command(update: Update, context):
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add')],
        [InlineKeyboardButton("Delete a book", callback_data='delete')],
        [InlineKeyboardButton("View Reading List", callback_data='view')]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text("Manage your reading list:", reply_markup=reply_markup)

async def button_callback(update: Update, context):
    query = update.callback_query
    await query.answer()

    if query.data == 'add':
        user_state[update.effective_user.id] = 'add'
        await query.edit_message_text(text="Please send the book title to add to your reading list.")

    elif query.data == 'delete':
        user_state[update.effective_user.id] = 'delete'
        await query.edit_message_text(text="Please send the book title to remove.")

    elif query.data == 'view':
        if reading_list:
            doc = Document()
            doc.add_heading('Reading List', 0)
            for book in reading_list:
                doc.add_paragraph(f"Title: {book['title']}\nPreview: {book.get('preview_link', 'N/A')}")
            doc_filename = "reading_list.docx"
            doc.save(doc_filename)

            with open(doc_filename, 'rb') as file:
                await query.message.reply_document(document=file, filename=doc_filename)
            os.remove(doc_filename)  # Clean up
        else:
            await query.edit_message_text(text="Your reading list is empty.")

async def handle_add_book(update: Update, context):
    # User is adding a book to the reading list
    book_name = update.message.text

    # Check if the book is already in the reading list to avoid duplicates
    for book in reading_list:
        if book['title'].lower() == book_name.lower():
            await update.message.reply_text(f"'{book_name}' is already in your reading list.")
            return

    preview_link = get_book_preview(book_name)  # Get the preview link from the API
    
    # Add the book to the reading list
    reading_list.append({'title': book_name, 'preview_link': preview_link})
    
    await update.message.reply_text(f"Added '{book_name}' to your reading list with preview link: {preview_link}")

async def handle_user_input(update: Update, context):
    user_id = update.effective_user.id

    # Check the user's current state and handle accordingly
    if user_state.get(user_id) == 'book':
        await handle_genre_search(update, context)

    elif user_state.get(user_id) == 'preview':
        await handle_preview(update, context)

    elif user_state.get(user_id) == 'add':
        # When the user is adding a book to the reading list
        await handle_add_book(update, context)

    elif user_state.get(user_id) == 'delete':
        # Handle deleting a book from the reading list
        await handle_delete_book(update, context)

async def handle_delete_book(update: Update, context):
    book_name = update.message.text
    global reading_list
    reading_list = [book for book in reading_list if book['title'].lower() != book_name.lower()]
    await update.message.reply_text(f"'{book_name}' has been removed from your reading list.")

def search_books_by_genre(genre):
    url = f"https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={GOOGLE_BOOKS_API_KEY}"
    response = requests.get(url)
    books = []
    if response.status_code == 200:
        data = response.json()
        for item in data.get('items', []):
            book_info = item['volumeInfo']
            books.append({
                'title': book_info.get('title', 'N/A'),
                'author': ', '.join(book_info.get('authors', [])),
                'description': book_info.get('description', 'No description available'),
                'published_year': book_info.get('publishedDate', 'N/A'),
                'language': book_info.get('language', 'N/A'),
                'preview_link': book_info.get('previewLink', 'No preview available')
            })
    return books

def get_book_preview(book_name):
    url = f"https://www.googleapis.com/books/v1/volumes?q={book_name}&key={GOOGLE_BOOKS_API_KEY}"
    response = requests.get(url)
    if response.status_code == 200:
        data = response.json()
        if 'items' in data and len(data['items']) > 0:
            return data['items'][0]['volumeInfo'].get('previewLink', None)
    return None

if __name__ == '__main__':
    application = ApplicationBuilder().token('7528399921:AAGtuW_ljSl60cOEr3PIENmcye5OGQXpiZQ').build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(CommandHandler("help", help_command))
    application.add_handler(CommandHandler("book", book_command))
    application.add_handler(CommandHandler("preview", preview_command))
    application.add_handler(CommandHandler("reading_list", reading_list_command))
    application.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_user_input))
    application.add_handler(CallbackQueryHandler(button_callback))

    application.run_polling()